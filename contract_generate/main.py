import os
import time
import zipfile

import docx
from PIL import Image
from pdf2image import convert_from_path
from win32com import client


def doc2pdf(doc_name, pdf_name):
    """
    :word文件转pdf
    :param doc_name word文件名称
    :param pdf_name 转换后pdf文件名称
    """
    try:
        word = client.DispatchEx("Word.Application")
        if os.path.exists(pdf_name):
            os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name, ReadOnly=1)
        worddoc.SaveAs(pdf_name, FileFormat=17)
        worddoc.Close()
        return pdf_name
    except Exception as e:
        print(e)


def demo_create(temp, file_name):
    doc = docx.Document(temp)

    # 每一段的编号、内容
    for i in range(len(doc.paragraphs)):
        print(str(i), doc.paragraphs[i].text)

        # 合同日期
        if i == 1:
            p = doc.paragraphs[i]
            run1 = p.add_run(str(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))))

        # 合同编号
        if i == 2:
            p = doc.paragraphs[i]
            run1 = p.add_run("330031118")

        # 联系人/电话
        if i == 6:
            style = doc.paragraphs[i].style
            text = doc.paragraphs[i].text
            text_list = text.split('：')
            # 手机号
            text_list[1] = text_list[1] + "：" + str(187878542154)
            text_list[0] = text_list[0] + "：" + str("xxxxxxxxxxxxx")
            p = doc.paragraphs[i].clear()
            doc.paragraphs[i].style = style
            run1 = p.add_run(str(text_list[0]) + str(text_list[1]))

        # 地址
        if i == 7:
            p = doc.paragraphs[i]
            run1 = p.add_run("xxxxx-xxxxxxxxxxxddd")

        # 买方
        if i == 69:
            p = doc.paragraphs[i]
            run1 = p.add_run("xxxxxxxxxxxxxxxxxx1")

    #     # 代表1
    #     if i == 72:
    #         p = doc.paragraphs[i]
    #         run1 = p.add_run("xxxxxx")
    #
    #     # 日期
    #     if i == 74:
    #         p = doc.paragraphs[i]
    #         run1 = p.add_run(str(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))))
    #
    #     # 代表2
    #     if i == 77:
    #         p = doc.paragraphs[i]
    #         run1 = p.add_run("xxxx")
    #
    #     # 日期
    #     if i == 79:
    #         p = doc.paragraphs[i]
    #         run1 = p.add_run(str(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))))
    #
    # # 表格
    # tbs = doc.tables
    # for tb in tbs:
    #     # 行
    #     for row in tb.rows:
    #         # 列
    #         i = 0
    #         for cell in row.cells:
    #             print(str(i), cell.text)
    #             i = i + 1
    #             # 也可以用下面方法
    #             '''text = ''
    #             for p in cell.paragraphs:
    #                 text += p.text
    #             print(text)'''

    doc.save(file_name)


# 图片添加水印
def pic_add_watermark(fn, watermark):
    file_name = fn.split('.')[0]
    list = os.listdir(os.getcwd())
    i = 0
    file_list = []
    # 寻找文件列表
    for file in list:
        if file_name in file and file.endswith(".png"):
            print(file)
            file_list.append(file)
            i = i + 1
    print(file_list)
    print(i)
    # 图片合成
    for file in file_list:
        background = Image.open(file)
        foreground = Image.open(watermark)
        background.paste(foreground, (background.size[0] - foreground.size[0] - 200, 150), foreground)
        if file == file_name + str(3) + ".png":
            background.paste(foreground, (background.size[0] - foreground.size[0] - 350, 1700), foreground)
            # background.show()
        background.save(file)

    im_list = []
    for each_image in range(i):
        try:
            img = Image.open(file_name + str((each_image + 2)) + ".png")
        except:
            continue
        if img.mode == "RGBA":
            img = img.convert('RGB')
            im_list.append(img)
        else:
            im_list.append(img)
    new_pdf = file_name + "_watermark.pdf"
    if not os.path.isfile(file_name + "_watermark.pdf"):  # 无文件时创建
        fd = open(new_pdf, mode="w", encoding="utf-8")
        fd.close()
    pdf = Image.open(file_list[0])
    # 保存pdf
    pdf.save(new_pdf, "PDF", resolution=100.0, save_all=True, append_images=im_list)
    # 删除图片
    for file in file_list:
        os.remove(file)
    return new_pdf


# pdf转png
def pdf_pic(fn):
    pages = convert_from_path(fn)
    file_name = fn.split('.')[0]
    for i in range(0, len(pages)):
        new_name = file_name + str((i + 1))
        pages[i].save(f"" + new_name + ".png", 'PNG')
        print("生成图片：{}".format(new_name))


if __name__ == '__main__':
    temp = ['free_temp.docx', 'pay_temp.docx']
    creat_temp = ['mark.pdf']
    for itm in temp:
        file_name = itm.split('.')[0] + str("_1") + ".docx"
        docx_name = file_name
        demo_create(itm, file_name)
        list = os.listdir(os.getcwd())
        file_name_pdf = itm.split('.')[0] + str("_1") + ".pdf"
        for file in list:
            if file not in temp and file.endswith(".docx"):
                doc2pdf(str(os.getcwd()) + "/" + file_name,
                        str(os.getcwd()) + "/" + file_name_pdf)
        list = os.listdir(os.getcwd())
        for file in list:
            if file not in creat_temp and file.endswith(".pdf"):
                # 打水印生成pdf
                file_name = file
                print(file)
                pdf_pic(file_name)
                new_pdf = pic_add_watermark(file_name, "a.png")

                creat_temp.append(new_pdf)
                os.remove(str(os.getcwd()) + "/" + file_name)
                os.remove(str(os.getcwd()) + "/" + docx_name)

    z = zipfile.ZipFile('contract.zip', 'w')
    for temp in creat_temp:
        if temp != 'mark.pdf':
            z.write(temp)
    z.close()
    # 删除
    for temp in creat_temp:
        if temp != 'mark.pdf':
            os.remove(temp)
